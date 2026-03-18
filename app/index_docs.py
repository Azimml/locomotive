
from __future__ import annotations

import os
import re
import subprocess
import sys
import tempfile

import chromadb
from docx import Document
from openai import OpenAI

DOCS_DIR = "/home/azim/Desktop/Perfect_RAG/docs"
CHROMA_PERSIST_DIR = "/home/azim/Desktop/Perfect_RAG/chroma_db"
COLLECTION_NAME = "repair_manuals"
EMBEDDING_MODEL = "text-embedding-3-small"
CHUNK_SIZE = 2400  

MODEL_MAP = {
    "2UZELR": "2UZ-EL(R)",
    "3ЭС5К": "3ЭС5К",
    "UZEL(R)": "UZ-EL(R)",
    "ВЛ80с": "ВЛ80С",
    "ТЭМ2": "ТЭМ2",
    "ТЭП70БС": "ТЭП70БС",
}


def extract_model_from_filename(filename: str) -> tuple[str, str]:
    base = os.path.splitext(filename)[0]
    parts = base.rsplit("_", 1)
    raw_model = parts[0] if len(parts) > 1 else base
    repair_type = parts[1] if len(parts) > 1 else "TXK2"
    model_name = MODEL_MAP.get(raw_model, raw_model)
    return model_name, repair_type


def convert_doc_to_docx(doc_path: str, output_dir: str) -> str:
    subprocess.run(
        [
            "libreoffice",
            "--headless",
            "--convert-to", "docx",
            "--outdir", output_dir,
            doc_path,
        ],
        check=True,
        capture_output=True,
    )
    basename = os.path.splitext(os.path.basename(doc_path))[0]
    return os.path.join(output_dir, f"{basename}.docx")


def extract_text_from_docx(docx_path: str) -> list[dict]:
    doc = Document(docx_path)
    paragraphs = []

    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue
        paragraphs.append({
            "text": text,
            "style": para.style.name if para.style else "Normal",
        })

    for table in doc.tables:
        for row in table.rows:
            row_texts = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if row_texts:
                paragraphs.append({
                    "text": " | ".join(row_texts),
                    "style": "TableRow",
                })

    return paragraphs


def chunk_paragraphs(
    paragraphs: list[dict],
    model_name: str,
    repair_type: str,
    source_file: str,
) -> list[dict]:
    heading_styles = {
        "Heading 1", "Heading 2", "Heading 3",
        "Heading1", "Heading2", "Heading3",
        "Title", "Subtitle",
    }

    sections: list[dict] = []
    current_section: dict = {"heading": "", "paragraphs": []}

    for para in paragraphs:
        is_heading = (
            para["style"] in heading_styles
            or re.match(r"^\d+\.\d*\s", para["text"])
            or re.match(r"^[IVXLC]+\.\s", para["text"])
        )

        if is_heading and current_section["paragraphs"]:
            sections.append(current_section)
            current_section = {"heading": para["text"], "paragraphs": []}
        elif is_heading:
            current_section["heading"] = para["text"]
        else:
            current_section["paragraphs"].append(para["text"])

    if current_section["paragraphs"]:
        sections.append(current_section)

    # Convert sections to chunks
    chunks = []
    chunk_index = 0

    for section in sections:
        section_text = "\n".join(section["paragraphs"])
        heading = section["heading"]

        if len(section_text) <= CHUNK_SIZE:
            full_text = f"{heading}\n\n{section_text}" if heading else section_text
            chunks.append({
                "text": full_text.strip(),
                "metadata": {
                    "locomotive_model": model_name,
                    "repair_type": repair_type,
                    "section_heading": heading,
                    "source_file": source_file,
                    "chunk_index": chunk_index,
                },
            })
            chunk_index += 1
        else:
            current_parts: list[str] = []
            current_len = 0

            for para_text in section["paragraphs"]:
                if current_len + len(para_text) > CHUNK_SIZE and current_parts:
                    full_text = f"{heading}\n\n" + "\n".join(current_parts)
                    chunks.append({
                        "text": full_text.strip(),
                        "metadata": {
                            "locomotive_model": model_name,
                            "repair_type": repair_type,
                            "section_heading": heading,
                            "source_file": source_file,
                            "chunk_index": chunk_index,
                        },
                    })
                    chunk_index += 1

                    overlap = current_parts[-2:] if len(current_parts) >= 2 else current_parts[-1:]
                    current_parts = list(overlap)
                    current_len = sum(len(p) for p in current_parts)

                current_parts.append(para_text)
                current_len += len(para_text)

            if current_parts:
                full_text = f"{heading}\n\n" + "\n".join(current_parts)
                chunks.append({
                    "text": full_text.strip(),
                    "metadata": {
                        "locomotive_model": model_name,
                        "repair_type": repair_type,
                        "section_heading": heading,
                        "source_file": source_file,
                        "chunk_index": chunk_index,
                    },
                })
                chunk_index += 1

    return chunks


def embed_texts(texts: list[str], openai_client: OpenAI) -> list[list[float]]:
    all_embeddings: list[list[float]] = []
    batch_size = 100

    for i in range(0, len(texts), batch_size):
        batch = texts[i : i + batch_size]
        response = openai_client.embeddings.create(
            model=EMBEDDING_MODEL,
            input=batch,
        )
        all_embeddings.extend([d.embedding for d in response.data])

    return all_embeddings


def main():
    api_key = os.environ.get("OPENAI_API_KEY")
    if not api_key:
        print("ERROR: OPENAI_API_KEY environment variable not set")
        sys.exit(1)

    openai_client = OpenAI(api_key=api_key)

    chroma_client = chromadb.PersistentClient(path=CHROMA_PERSIST_DIR)

    try:
        chroma_client.delete_collection(COLLECTION_NAME)
        print(f"Deleted existing collection '{COLLECTION_NAME}'")
    except Exception:
        pass

    collection = chroma_client.get_or_create_collection(
        name=COLLECTION_NAME,
        metadata={"hnsw:space": "cosine"},
    )

    doc_files = [f for f in os.listdir(DOCS_DIR) if f.endswith(".doc")]
    print(f"Found {len(doc_files)} .doc files\n")

    all_chunks: list[dict] = []

    with tempfile.TemporaryDirectory() as tmp_dir:
        for doc_file in doc_files:
            doc_path = os.path.join(DOCS_DIR, doc_file)
            model_name, repair_type = extract_model_from_filename(doc_file)

            print(f"Processing: {doc_file}")
            print(f"  Model: {model_name}, Repair type: {repair_type}")

            try:
                docx_path = convert_doc_to_docx(doc_path, tmp_dir)
            except subprocess.CalledProcessError as e:
                print(f"  ERROR converting {doc_file}: {e}")
                continue

            paragraphs = extract_text_from_docx(docx_path)
            print(f"  Extracted {len(paragraphs)} paragraphs")

            chunks = chunk_paragraphs(paragraphs, model_name, repair_type, doc_file)
            print(f"  Created {len(chunks)} chunks")
            all_chunks.extend(chunks)

    print(f"\nTotal chunks: {len(all_chunks)}")

    if not all_chunks:
        print("No chunks to index. Exiting.")
        return

    print("Embedding chunks...")
    texts = [c["text"] for c in all_chunks]
    embeddings = embed_texts(texts, openai_client)

    print("Storing in ChromaDB...")
    ids = [f"chunk_{i}" for i in range(len(all_chunks))]
    metadatas = [c["metadata"] for c in all_chunks]

    batch_size = 100
    for i in range(0, len(all_chunks), batch_size):
        end = min(i + batch_size, len(all_chunks))
        collection.add(
            ids=ids[i:end],
            documents=texts[i:end],
            embeddings=embeddings[i:end],
            metadatas=metadatas[i:end],
        )

    print(f"\nDone! Indexed {len(all_chunks)} chunks into ChromaDB")
    print(f"Persisted at: {CHROMA_PERSIST_DIR}\n")

    # Print summary
    for doc_file in doc_files:
        model_name, _ = extract_model_from_filename(doc_file)
        count = sum(1 for c in all_chunks if c["metadata"]["source_file"] == doc_file)
        print(f"  {model_name}: {count} chunks")


if __name__ == "__main__":
    main()
